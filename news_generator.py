import requests
import json
import datetime
import time
import os
import math
from bs4 import BeautifulSoup
from groq import Groq
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from urllib.parse import urlparse
from io import BytesIO
from dotenv import load_dotenv
from html2image import Html2Image

load_dotenv()

hti = Html2Image()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))
document = Document()
headers_requests = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:77.0) Gecko/20100101 Firefox/77.0"}

classements = ["https://sport.detik.com/sepakbola/klasemen/liga-indonesia", "https://sport.detik.com/sepakbola/klasemen/liga-spanyol", "https://sport.detik.com/sepakbola/klasemen/liga-inggris", "https://id-mpl.com/"]

news_structure = {"src": "", "page": {"width": 8.27, "height": 11.69, "mtop": 0.5, "mbot": 0.5, "mlef": 0.5, "mrig": 0.5}, "title": {"font": "Times New Roman", "size": 22, "bold": True, "content": ""}, "quotes": {"font": "Times New Roman", "size": 10, "italic": True, "content": ""}, "image": {"height": 2.36, "src": ""}, "paragraphs": [{"font": "Times New Roman", "size": 11, "content": ""}]}

system_prompt = f"Buat berita yang di kirimkan oleh user menjadi berita yang menarik dan bermodel paragraf, gunakan format json di setiap style font, gunakan struktur seperti berikut {str(news_structure)}, pastikan isi berita dapat pas pada ukuran kertas yang diberikan (Inch), ukuran font (pt), gunakan bahasa Indonesia, jika berita yang diberikan oleh user kosong maka tuliskan gagal di setiap content, berikan dateline di awal berita yang mengandung tempat dan sumber berita, jangan berikan tulisan copyright, buat 6 paragraf, buat masing masing paragraf berisi 5 baris, gunakan gaya bahasa seperti artikel ilmiah populer"

def generate_news(news):
  for i in range(3):
    try:
      generated_news = client.chat.completions.create(messages=[
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": news}
      ], model="deepseek-r1-distill-llama-70b", reasoning_format="parsed", response_format={"type": "json_object"}, stream=False)
      result = generated_news.choices[0].message.content
      return json.loads(result)
    except:
      print(f"News failed to generate. Attemps: {i+1}")
      time.sleep(0.5)

def news_filter(url, element):
  domain = urlparse(url).netloc

  inner_news = element.find_all("p", {"class": None})
  if (domain.endswith("cnnindonesia.com")):
    image = element.find("img", {"class": "w-full"})
  elif (domain.endswith("kompas.com")):
    image_wrap = element.find("div", {"class": "photo__wrap"})
    image = None
    if (image_wrap is not None):
      image = image_wrap.find("img")
  elif (domain.endswith("detik.com")):
    image = element.find("img", {"class": "img-zoomin"})
  elif (domain.endswith("tribunnews.com")):
    image = element.find("img", {"class": "imgfull"})
  elif (domain.endswith("kumparan.com")):
    inner_news = element.find_all("span", {"class": "Textweb__StyledText-sc-1ed9ao-0"})
    image = element.find("img", {"class": "ImageLoaderweb__StyledImage-sc-zranhd-0"})
  elif (domain.endswith("bola.com")):
    image = element.find("img", {"class": "read-page--photo-gallery--item__picture-lazyload"})
  elif (domain.endswith("kincir.com")):
    image = element.find("img", {"class": "wp-post-image"})
  
  return {"inner_news": inner_news, "image": image}


def get_news(url):
  res_arr = []
  image = None
  for i in range(3):
    try:
      response = requests.get(url=url, headers=headers_requests)
      element = BeautifulSoup(response.content, "html.parser")
      filtered = news_filter(url=url, element=element)
      image = filtered["image"]
      for e in filtered["inner_news"]:
        text = e.text
        if text != "":
          res_arr.append(text)
      if res_arr == []:
        print(f"Getting news failed. Attemp: {i+1}")
        time.sleep(0.5)
      else:
        break
    except Exception as e:
      print(e)
      print(f"Getting news failed. Attemp: {i+1}")
      time.sleep(0.5)
  return {"news": " ".join(res_arr), "img": image}

def create_page(data, img, news_length, news_index):
  title = document.add_paragraph()
  title_format = title.paragraph_format
  title_format.space_after = Pt(0)
  title_font = title.add_run(data["title"]["content"]).font
  title_font.name = data["title"]["font"]
  title_font.size = Pt(data["title"]["size"])
  title_font.bold = data["title"]["bold"]

  quotes = document.add_paragraph()
  quotes_font = quotes.add_run(data["quotes"]["content"]).font
  quotes_font.size = Pt(data["quotes"]["size"])
  quotes_font.italic = data["quotes"]["italic"]
  quotes_font.name = data["quotes"]["font"]

  try:
    if (img["src"] == "placeholder"):
      document.add_picture("./assets/placeholder.png", height=Inches(img["height"]))
    else:
      response = requests.get(url=img["src"], headers=headers_requests)
      image_stream = BytesIO(response.content)
      document.add_picture(image_stream, height=Inches(img["height"]))
  except:
      document.add_picture("./assets/placeholder.png", height=Inches(img["height"]))

  last_paragraph = document.paragraphs[-1]
  last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

  for p in data["paragraphs"]:
    paragraph = document.add_paragraph()
    paragraph_font = paragraph.add_run(p["content"]).font
    paragraph_font.name = p["font"]
    paragraph_font.size = Pt(p["size"])
  if news_index+1 < news_length:
    document.add_page_break()

def get_classement(url):
  domain = urlparse(url).netloc
  response = requests.get(url)
  element = BeautifulSoup(response.content, "html.parser")
  hti = Html2Image()
  hti.output_path = "./assets/"
  style_file = element.find("link", {"rel": "stylesheet"})
  style_inline = element.find("style")
  table = ""
  if domain.endswith("detik.com"):
    custom_style = "<style>.new_klasemen .klasemen-table > span:nth-of-type(3) { width: 40px; justify-content: center; } .new_klasemen .klasemen-table > span:nth-of-type(2) { width: 350px; justify-content: flex-start; }</style>"
    rows = element.find_all("div", {"class": "klasemen-table"})[:16]
    for row in rows:
      all_spans = row.find_all("span", recursive=False)
      row_class = " ".join(row["class"])
      spans_without_more = "".join(str(x) for x in all_spans[1:])
      new_row = "".join(f"<div class=\"{row_class}\">{spans_without_more}</div>")
      table += new_row
    hti.screenshot(html_str=f"{str(style_file)}{custom_style}<div class=\"new_klasemen\">{table}</div>", size=(895,770))
  elif domain.endswith("mpl.com"):
    custom_style = "<style>table { width: 1000px } thead { background-color: black; color: white !important; } tbody > tr:nth-child(2n) { background-color: #eaeaea } tr * { text-align: center; align-items: center }</style>"
    table = element.find("div", {"id": "standing-regular-season"})
    hti.screenshot(f"{str(style_file)}{str(style_inline)}{custom_style}{table}", size=(1000, 470))

def create_document(data):
  section = document.sections[0]
  section.page_width = Inches(data["page"]["width"])
  section.page_height = Inches(data["page"]["height"])
  section.top_margin = Inches(data["page"]["mtop"])
  section.bottom_margin = Inches(data["page"]["mbot"])
  section.left_margin = Inches(data["page"]["mlef"])
  section.right_margin = Inches(data["page"]["mrig"])

  document.add_page_break()

  for classement in classements:
    get_classement(classement)
    time.sleep(0.3)
    document.add_picture("./assets/screenshot.png", width=Inches(4.3))

  week = math.floor(datetime.datetime.now().day/7)
  month, year = datetime.datetime.now().month, datetime.datetime.now().year
  if week == 0 and datetime.datetime.now().weekday() > 6:
    week = 4
    month -= 1
    if month == 0:
      year -=1
  
  formatted_date = datetime.datetime(month=month, year=year, day=1)

  edition = f"{formatted_date.strftime("%B %Y")} ({str(week)})"
  section.footer.add_paragraph(f"Santri Update {edition}")
  
  document.save(f"sans {edition}.docx")

def make_news(url_arr):
  news_result = []
  for index, url in enumerate(url_arr):
    for i in range(3):
      try:
        scrape_result = get_news(url)
        news = generate_news(scrape_result["news"])
        if scrape_result["img"] is not None:
          news["image"]["src"] = scrape_result["img"]["src"]
        else:
          news["image"]["src"] = "placeholder"
        news["src"] = urlparse(url).netloc
        news_result.append(news)
        create_page(data=news, img=news["image"], news_length=len(url_arr), news_index=index)
        print(f"News ({index+1}/{len(url_arr)}) created: {news["title"]["content"]}")
        break
      except Exception as e:
        print(e)
        print(f"Creating news ({index+1}/{len(url_arr)}) failed. Retrying ({i+1})")
        time.sleep(0.3)
    time.sleep(0.3)
  create_document(news_result[0])