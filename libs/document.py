import datetime, math, os
from libs import scraper
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from html2image import Html2Image

hti = Html2Image()
document = Document()

def create_page(data, img, news_length, news_index):
  title = document.add_paragraph()
  title_format = title.paragraph_format
  title_format.space_after = Pt(0)
  title_font = title.add_run(data["title"]["content"]).font
  title_font.name = data["title"]["font"]
  title_font.size = Pt(data["title"]["size"])
  title_font.bold = data["title"]["bold"]

  quotes = document.add_paragraph()
  quotes_font = quotes.add_run(data["quotes"]["content"]).font
  quotes_font.size = Pt(data["quotes"]["size"])
  quotes_font.italic = data["quotes"]["italic"]
  quotes_font.name = data["quotes"]["font"]

  try:
    document.add_picture(img["content"], height=Inches(img["height"]))
  except Exception as e:
    print(e)

  last_paragraph = document.paragraphs[-1]
  last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

  for p in data["paragraphs"]:
    paragraph = document.add_paragraph()
    paragraph_font = paragraph.add_run(p["content"]).font
    paragraph_font.name = p["font"]
    paragraph_font.size = Pt(p["size"])
  if news_index+1 < news_length:
    document.add_page_break()



def create_document(data, classement_list, use_footer, use_classements):
  section = document.sections[0]
  section.page_width = Inches(data["page"]["width"])
  section.page_height = Inches(data["page"]["height"])
  section.top_margin = Inches(data["page"]["mtop"])
  section.bottom_margin = Inches(data["page"]["mbot"])
  section.left_margin = Inches(data["page"]["mlef"])
  section.right_margin = Inches(data["page"]["mrig"])

  document.add_page_break()

  if use_classements:
    for classement in classement_list:
      scraper.get_classement(classement)
      document.add_picture(".cache/screenshot.png", width=Inches(4.3))

  datetime_now = datetime.datetime.now()
  first_weekday = datetime.datetime(year=datetime_now.year, month=datetime_now.month, day=1).weekday()

  offset_date = datetime_now.day + first_weekday - 1
  week = math.floor(offset_date/7)
  month, year = datetime_now.month, datetime_now.year
  if week == 0:
    week = 4
    month -= 1
    if month == 0:
      month = 12
      year -=1
  
  formatted_date = datetime.datetime(month=month, year=year, day=1)

  edition = f"{formatted_date.strftime("%B %Y")} ({str(week)})"
  if use_footer: section.footer.add_paragraph(f"Santri Update {edition}")
  
  out_dir = f"SANS/{formatted_date.strftime("%Y/%B")}"
  os.makedirs(out_dir, exist_ok=True)

  document.save(f"{out_dir}/sans {edition}.docx")

